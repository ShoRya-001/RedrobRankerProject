from __future__ import annotations

import html
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

import streamlit as st

# Import the dependency-free ranker from the repository root.
REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from rank import iter_candidates, rank_candidates, write_submission  # noqa: E402
from validate_submission import validate_submission  # noqa: E402


st.set_page_config(
    page_title="Redrob Candidate Ranker",
    layout="wide",
    initial_sidebar_state="expanded",
)


CUSTOM_CSS = """
<style>
:root {
  color-scheme: light;
  --accent: #1DB954;
  --accent-hover: #22D660;
  --accent-soft: rgba(29, 185, 84, 0.12);
  --accent-border: rgba(29, 185, 84, 0.28);
  --info: #2563EB;
  --info-soft: rgba(37, 99, 235, 0.12);
  --warning: #F59E0B;
  --warning-soft: rgba(245, 158, 11, 0.14);
  --error: #EF4444;
  --error-soft: rgba(239, 68, 68, 0.12);

  --bg: #F8F9FA;
  --bg-2: #F2F3F5;
  --surface: #FFFFFF;
  --surface-2: #F7F8FA;
  --surface-hover: #EEF1F4;
  --glass: rgba(255, 255, 255, 0.82);
  --border: #E5E7EB;
  --border-soft: rgba(17, 24, 39, 0.08);
  --text: #111827;
  --text-2: #4B5563;
  --muted: #6B7280;
  --sidebar-bg: #FFFFFF;
  --sidebar-text: #111827;
  --sidebar-muted: #4B5563;
  --shadow: 0 18px 55px rgba(17, 24, 39, 0.10);
  --shadow-2: 0 10px 28px rgba(17, 24, 39, 0.08);
  --code-bg: #111827;
  --code-text: #ECFDF5;
}

@media (prefers-color-scheme: dark) {
  :root {
    color-scheme: dark;
    --bg: #121212;
    --bg-2: #181818;
    --surface: #202020;
    --surface-2: #181818;
    --surface-hover: #282828;
    --glass: rgba(32, 32, 32, 0.82);
    --border: rgba(255,255,255,0.08);
    --border-soft: rgba(255,255,255,0.08);
    --text: #FFFFFF;
    --text-2: #B3B3B3;
    --muted: #8A8A8A;
    --sidebar-bg: #000000;
    --sidebar-text: #FFFFFF;
    --sidebar-muted: #B3B3B3;
    --shadow: 0 22px 70px rgba(0,0,0,0.45);
    --shadow-2: 0 12px 32px rgba(0,0,0,0.35);
    --code-bg: #050505;
    --code-text: #E5E7EB;
  }
}

* {
  transition: background-color 220ms ease, border-color 220ms ease, box-shadow 220ms ease, transform 180ms ease, color 160ms ease, opacity 180ms ease;
}

.stApp {
  background:
    radial-gradient(circle at 12% 8%, rgba(29, 185, 84, 0.13), transparent 28rem),
    radial-gradient(circle at 92% 0%, rgba(37, 99, 235, 0.08), transparent 24rem),
    radial-gradient(circle at 50% 100%, rgba(29, 185, 84, 0.07), transparent 30rem),
    linear-gradient(180deg, var(--bg) 0%, var(--bg-2) 100%);
  color: var(--text);
  font-family: Inter, Manrope, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "SF Pro Display", "Segoe UI", sans-serif;
}

.block-container {
  width: 92%;
  max-width: 1480px;
  padding-top: 1.15rem;
  padding-bottom: 2.5rem;
}

#MainMenu, footer, header[data-testid="stHeader"] {
  visibility: hidden;
}

[data-testid="stSidebar"] {
  background: var(--sidebar-bg);
  border-right: 1px solid var(--border);
  box-shadow: 10px 0 34px rgba(0,0,0,0.08);
}

[data-testid="stSidebar"] * {
  color: var(--sidebar-text) !important;
}

[data-testid="stSidebar"] small,
[data-testid="stSidebar"] .stCaptionContainer,
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p {
  color: var(--sidebar-muted) !important;
}

section[data-testid="stSidebar"] div[role="radiogroup"] label {
  min-height: 46px;
  padding: 0.65rem 0.75rem;
  border-radius: 16px;
  margin-bottom: 0.34rem;
  background: transparent;
  border: 1px solid transparent;
}

section[data-testid="stSidebar"] div[role="radiogroup"] label:hover {
  background: var(--surface-hover);
  border-color: var(--border);
  transform: translateX(2px);
}

section[data-testid="stSidebar"] div[role="radiogroup"] label:has(input:checked) {
  background: var(--accent-soft);
  border-color: var(--accent-border);
}

h1, h2, h3, h4, h5, h6, p, li, label, span, div {
  color: var(--text);
}

p, li {
  color: var(--text-2);
  line-height: 1.68;
}

h1, h2, h3 {
  letter-spacing: -0.045em;
}

a { color: var(--accent) !important; }

.sidebar-brand {
  display: flex;
  align-items: center;
  gap: 12px;
  margin: 8px 0 10px;
}

.sidebar-logo {
  width: 44px;
  height: 44px;
  border-radius: 14px;
  display: grid;
  place-items: center;
  background: var(--accent);
  color: #000000;
  box-shadow: 0 12px 32px rgba(29,185,84,0.28);
}

.sidebar-title {
  font-size: 1.15rem;
  line-height: 1.1;
  font-weight: 950;
  letter-spacing: -0.04em;
}

.sidebar-copy {
  margin: 0 0 16px;
  color: var(--sidebar-muted) !important;
  font-size: 0.91rem;
}

.check-progress-label {
  display: flex;
  justify-content: space-between;
  align-items: center;
  gap: 8px;
  margin: 8px 0 8px;
  color: var(--sidebar-muted) !important;
  font-size: 0.86rem;
  font-weight: 800;
}

.check-progress-track {
  height: 8px;
  width: 100%;
  border-radius: 999px;
  overflow: hidden;
  background: var(--surface-hover);
  border: 1px solid var(--border);
  margin-bottom: 14px;
}

.check-progress-fill {
  height: 100%;
  border-radius: 999px;
  background: linear-gradient(90deg, var(--accent), #7DE39B);
  animation: progressIn 420ms ease both;
}

.checklist-wrap {
  display: grid;
  gap: 10px;
}

.check-item {
  display: grid;
  grid-template-columns: 32px 1fr;
  gap: 11px;
  align-items: center;
  min-height: 40px;
  color: var(--sidebar-muted);
  font-size: 0.92rem;
}

.check-dot {
  width: 27px;
  height: 27px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  border: 2px solid color-mix(in srgb, var(--sidebar-muted) 58%, transparent);
  color: transparent;
  font-weight: 950;
}

.check-dot.done {
  border-color: var(--accent);
  background: var(--accent);
  color: #000;
  box-shadow: 0 10px 24px rgba(29,185,84,0.28);
  animation: tickPop 260ms ease both;
}

.check-label.done {
  color: var(--sidebar-text) !important;
  font-weight: 850;
}

.spotify-shell {
  position: sticky;
  top: 0;
  z-index: 20;
  margin-bottom: 14px;
  border: 1px solid var(--border);
  border-radius: 20px;
  background: color-mix(in srgb, var(--glass) 88%, transparent);
  backdrop-filter: blur(18px);
  box-shadow: var(--shadow-2);
  padding: 12px 16px;
}

.topbar {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}

.brand-lockup {
  display: flex;
  align-items: center;
  gap: 11px;
}

.brand-icon {
  width: 40px;
  height: 40px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  background: var(--accent);
  color: #000;
  box-shadow: 0 12px 32px rgba(29,185,84,0.28);
}

.brand-title {
  font-weight: 950;
  letter-spacing: -0.05em;
  font-size: 1.05rem;
  color: var(--text);
}

.brand-subtitle {
  font-size: 0.80rem;
  color: var(--muted);
  font-weight: 750;
}

.topbar-badges {
  display: flex;
  flex-wrap: wrap;
  justify-content: flex-end;
  gap: 8px;
}

.badge {
  display: inline-flex;
  align-items: center;
  gap: 7px;
  border-radius: 999px;
  padding: 7px 11px;
  border: 1px solid var(--border);
  color: var(--text-2);
  background: var(--surface-2);
  font-size: 0.76rem;
  font-weight: 850;
  white-space: nowrap;
}

.badge-green {
  background: var(--accent-soft);
  border-color: var(--accent-border);
  color: var(--text);
}

.hero-card {
  position: relative;
  overflow: hidden;
  border-radius: 22px;
  border: 1px solid var(--border);
  background:
    radial-gradient(circle at 88% 4%, rgba(29,185,84,0.18), transparent 18rem),
    linear-gradient(135deg, var(--surface) 0%, var(--surface-2) 100%);
  box-shadow: var(--shadow);
  padding: clamp(1.15rem, 3vw, 1.65rem);
  margin-bottom: 16px;
  animation: fadeSlide 280ms ease both;
}

.hero-card:before {
  content: "";
  position: absolute;
  inset: 0;
  background: linear-gradient(90deg, rgba(29,185,84,0.07), transparent 42%);
  pointer-events: none;
}

.hero-content {
  position: relative;
  z-index: 1;
}

.hero-title {
  max-width: 850px;
  margin: 0;
  color: var(--text);
  font-size: clamp(2rem, 4.2vw, 4.6rem);
  line-height: 0.98;
  letter-spacing: -0.075em;
  font-weight: 980;
}

.gradient-text {
  background: linear-gradient(90deg, var(--accent), #7DE39B 60%, var(--text));
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
}

.hero-subtitle {
  max-width: 760px;
  margin-top: 12px;
  margin-bottom: 0;
  color: var(--text-2);
  font-size: 1.02rem;
  line-height: 1.58;
}

.quick-card, .glass-card, .method-card, .success-card, .empty-state, .step-card {
  border-radius: 20px;
  border: 1px solid var(--border);
  background: var(--glass);
  box-shadow: var(--shadow-2);
  padding: 18px;
  height: 100%;
  backdrop-filter: blur(14px);
  animation: fadeSlide 260ms ease both;
}

.quick-card:hover, .glass-card:hover, .method-card:hover, .success-card:hover, .step-card:hover {
  background: var(--surface-hover);
  transform: translateY(-2px) scale(1.004);
  box-shadow: var(--shadow);
}

.step-grid {
  display: grid;
  grid-template-columns: repeat(5, minmax(140px, 1fr));
  gap: 12px;
  margin: 0 0 18px;
}

.step-card {
  min-height: 104px;
}

.step-card.completed {
  border-color: var(--accent-border);
  background: var(--accent-soft);
}

.step-card.active {
  border-color: var(--accent-border);
  box-shadow: 0 0 0 3px rgba(29,185,84,0.10), var(--shadow-2);
}

.step-kicker {
  color: var(--muted);
  font-size: 0.70rem;
  font-weight: 950;
  letter-spacing: 0.14em;
  text-transform: uppercase;
}

.step-title {
  margin-top: 8px;
  color: var(--text);
  font-size: 0.96rem;
  font-weight: 930;
  line-height: 1.2;
}

.step-check {
  float: right;
  width: 24px;
  height: 24px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  background: var(--accent);
  color: #000;
  font-weight: 950;
  animation: tickPop 260ms ease both;
}

.step-pill {
  display: inline-flex;
  align-items: center;
  gap: 7px;
  padding: 7px 11px;
  border-radius: 999px;
  color: var(--text);
  background: var(--accent-soft);
  border: 1px solid var(--accent-border);
  font-size: 0.75rem;
  font-weight: 900;
  letter-spacing: 0.04em;
  margin-bottom: 12px;
}

.card-title {
  margin: 0 0 8px 0;
  color: var(--text);
  font-size: 1.08rem;
  font-weight: 950;
  letter-spacing: -0.035em;
}

.card-copy {
  margin: 0;
  color: var(--text-2);
  font-size: 0.92rem;
}

.status-card {
  border-radius: 18px;
  border: 1px solid var(--border);
  background: var(--glass);
  box-shadow: var(--shadow-2);
  padding: 16px;
  animation: fadeSlide 300ms ease both;
}

.status-icon {
  width: 28px;
  height: 28px;
  margin-bottom: 10px;
  color: var(--accent);
}

.status-label {
  color: var(--muted);
  font-size: 0.69rem;
  font-weight: 900;
  letter-spacing: 0.16em;
  text-transform: uppercase;
}

.status-value {
  color: var(--text);
  margin-top: 3px;
  font-size: 1.28rem;
  font-weight: 950;
  letter-spacing: -0.04em;
}

.stButton > button {
  min-height: 46px !important;
  border: 0 !important;
  border-radius: 999px !important;
  padding: 0.72rem 1.25rem !important;
  font-weight: 950 !important;
  color: #000000 !important;
  background: var(--accent) !important;
  box-shadow: 0 14px 34px rgba(29,185,84,0.24) !important;
}

.stButton > button:hover {
  background: var(--accent-hover) !important;
  color: #000000 !important;
  transform: translateY(-2px);
  box-shadow: 0 18px 44px rgba(29,185,84,0.32) !important;
}

.stButton > button:disabled {
  background: var(--surface-hover) !important;
  color: var(--muted) !important;
  box-shadow: none !important;
  cursor: not-allowed !important;
}

.stButton > button:focus-visible,
.stDownloadButton > button:focus-visible,
button:focus-visible,
input:focus-visible,
textarea:focus-visible {
  outline: 3px solid rgba(29,185,84,0.42) !important;
  outline-offset: 2px !important;
}

.stDownloadButton > button {
  min-height: 46px !important;
  border-radius: 999px !important;
  font-weight: 950 !important;
  border: 1px solid var(--accent-border) !important;
  color: var(--text) !important;
  background: var(--accent-soft) !important;
}

/* UPDATED: uploader visual polish only */
[data-testid="stFileUploader"] section {
  border: 2px dashed rgba(29, 185, 84, 0.72) !important;
  border-radius: 20px !important;
  background: rgba(29, 185, 84, 0.12) !important;
  padding: 18px !important;
  box-shadow: inset 0 0 0 1px rgba(29, 185, 84, 0.14), 0 10px 28px rgba(0, 0, 0, 0.16);
}

[data-testid="stFileUploader"] section:hover {
  border-color: #1DB954 !important;
  background: rgba(29, 185, 84, 0.18) !important;
  box-shadow: 0 0 0 3px rgba(29, 185, 84, 0.20), 0 16px 36px rgba(0, 0, 0, 0.20);
}

[data-testid="stFileUploader"] button {
  background: #1DB954 !important;
  color: #0B0B0B !important;
  border: 1px solid rgba(29, 185, 84, 0.78) !important;
  border-radius: 12px !important;
  min-height: 48px !important;
  padding: 0.68rem 1.25rem !important;
  font-weight: 900 !important;
  font-size: 0.95rem !important;
  box-shadow: 0 10px 24px rgba(29, 185, 84, 0.28) !important;
  cursor: pointer !important;
}

[data-testid="stFileUploader"] button:hover {
  background: #23D564 !important;
  color: #060606 !important;
  border-color: #23D564 !important;
  box-shadow: 0 14px 30px rgba(29, 185, 84, 0.34) !important;
}

[data-testid="stFileUploader"] small,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] p {
  color: var(--text-2) !important;
}

textarea, input {
  border-radius: 14px !important;
  color: var(--text) !important;
  background: var(--surface) !important;
  border-color: var(--border) !important;
}

textarea::placeholder, input::placeholder {
  color: var(--muted) !important;
  opacity: 1 !important;
}

textarea:focus, input:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(29,185,84,0.16) !important;
}

[data-testid="stMetric"] {
  border-radius: 18px;
  border: 1px solid var(--border);
  background: var(--glass);
  box-shadow: var(--shadow-2);
  padding: 16px;
}

[data-testid="stMetricLabel"] p {
  color: var(--text-2) !important;
  font-weight: 800;
}

[data-testid="stMetricValue"] {
  color: var(--text) !important;
  font-weight: 950 !important;
  letter-spacing: -0.05em;
}

.stAlert {
  border-radius: 16px !important;
  border: 1px solid var(--border) !important;
}

pre, code {
  background: var(--code-bg) !important;
  color: var(--code-text) !important;
  border-radius: 14px !important;
}

.success-card {
  display: grid;
  grid-template-columns: 42px 1fr;
  gap: 14px;
  align-items: start;
  margin-top: 12px;
  border-color: var(--accent-border);
  background: var(--accent-soft);
}

.success-icon {
  width: 38px;
  height: 38px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  color: #000;
  background: var(--accent);
  font-weight: 950;
  animation: tickPop 260ms ease both;
}

.success-title {
  font-weight: 950;
  color: var(--text);
  margin-bottom: 4px;
}

.success-meta {
  color: var(--text-2);
  font-size: 0.88rem;
  line-height: 1.55;
}

.sticky-run-panel {
  position: sticky;
  bottom: 18px;
  z-index: 15;
  border: 1px solid var(--border);
  border-radius: 20px;
  background: color-mix(in srgb, var(--glass) 92%, transparent);
  box-shadow: var(--shadow);
  backdrop-filter: blur(18px);
  padding: 14px;
}

.results-table-wrap {
  margin-top: 18px;
  overflow-x: auto;
  border-radius: 20px;
  border: 1px solid var(--border);
  background: var(--surface);
  box-shadow: var(--shadow-2);
}

.results-table {
  width: 100%;
  border-collapse: separate;
  border-spacing: 0;
  min-width: 980px;
}

.results-table thead th {
  position: sticky;
  top: 0;
  z-index: 1;
  background: var(--surface-2);
  color: var(--muted);
  font-size: 0.72rem;
  text-align: left;
  text-transform: uppercase;
  letter-spacing: 0.12em;
  padding: 14px 16px;
  border-bottom: 1px solid var(--border);
}

.results-table tbody tr:nth-child(even) {
  background: color-mix(in srgb, var(--surface-2) 50%, transparent);
}

.results-table tbody tr:hover {
  background: var(--surface-hover);
}

.results-table tbody tr.top-three {
  background: color-mix(in srgb, var(--accent-soft) 58%, transparent);
}

.results-table tbody tr.rank-one {
  box-shadow: inset 4px 0 0 #D4AF37;
}

.results-table td {
  color: var(--text-2);
  padding: 14px 16px;
  border-bottom: 1px solid var(--border-soft);
  vertical-align: top;
}

.results-table .candidate-id {
  color: var(--text);
  font-weight: 900;
}

.rank-chip {
  display: inline-flex;
  width: 34px;
  height: 34px;
  align-items: center;
  justify-content: center;
  border-radius: 999px;
  font-weight: 950;
  color: #000;
  background: var(--accent);
}

.rank-chip.gold {
  background: linear-gradient(135deg, #FDE68A, #D4AF37);
}

.score-pill {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 72px;
  border-radius: 999px;
  padding: 6px 10px;
  font-weight: 950;
}

.score-high { background: rgba(29,185,84,0.18); color: var(--accent); }
.score-mid { background: var(--warning-soft); color: var(--warning); }
.score-low { background: rgba(107,114,128,0.14); color: var(--text-2); }

.match-track {
  height: 10px;
  min-width: 140px;
  border-radius: 999px;
  background: var(--surface-hover);
  overflow: hidden;
  margin-top: 8px;
}

.match-fill {
  height: 100%;
  border-radius: 999px;
  background: linear-gradient(90deg, var(--accent), #7DE39B);
  animation: progressIn 580ms ease both;
}

.skill-chip {
  display: inline-flex;
  align-items: center;
  border-radius: 999px;
  padding: 5px 9px;
  margin: 0 5px 5px 0;
  background: var(--info-soft);
  color: var(--text);
  font-size: 0.78rem;
  font-weight: 800;
}

.preview-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 14px;
  margin: 18px 0;
}

.preview-card {
  border-radius: 20px;
  border: 1px solid var(--border);
  background: var(--glass);
  box-shadow: var(--shadow-2);
  padding: 18px;
  min-height: 220px;
  animation: fadeSlide 300ms ease both;
}

.preview-card:hover {
  transform: translateY(-2px) scale(1.004);
  background: var(--surface-hover);
}

.preview-card.rank-one {
  border-color: rgba(212, 175, 55, 0.62);
  box-shadow: 0 18px 56px rgba(212,175,55,0.14), var(--shadow-2);
}

.preview-rank {
  color: var(--muted);
  font-size: 0.76rem;
  font-weight: 950;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.preview-name {
  margin-top: 10px;
  color: var(--text);
  font-size: 1.18rem;
  font-weight: 950;
  letter-spacing: -0.04em;
}

.preview-score {
  margin: 12px 0;
  font-size: 2rem;
  line-height: 1;
  font-weight: 980;
  letter-spacing: -0.06em;
  color: var(--accent);
}

.empty-state {
  min-height: 320px;
  display: grid;
  place-items: center;
  text-align: center;
}

.empty-illustration {
  width: 94px;
  height: 94px;
  border-radius: 26px;
  display: grid;
  place-items: center;
  margin: 0 auto 18px;
  background: var(--accent-soft);
  color: var(--accent);
  border: 1px solid var(--accent-border);
}

.footer-card {
  margin-top: 28px;
  border-radius: 20px;
  border: 1px solid var(--border);
  background: var(--glass);
  box-shadow: var(--shadow-2);
  padding: 16px 18px;
  display: flex;
  flex-wrap: wrap;
  justify-content: space-between;
  gap: 12px;
  color: var(--text-2);
}

.footer-links {
  display: flex;
  gap: 14px;
  flex-wrap: wrap;
  font-weight: 850;
}

hr { border-color: var(--border) !important; }

@keyframes tickPop {
  0% { transform: scale(0.72); opacity: 0.55; }
  70% { transform: scale(1.12); opacity: 1; }
  100% { transform: scale(1); opacity: 1; }
}

@keyframes fadeSlide {
  from { opacity: 0; transform: translateY(8px); }
  to { opacity: 1; transform: translateY(0); }
}

@keyframes progressIn {
  from { width: 0; }
}

@media (max-width: 980px) {
  .step-grid, .preview-grid { grid-template-columns: 1fr; }
}

@media (max-width: 760px) {
  .topbar { align-items: flex-start; flex-direction: column; }
  .topbar-badges { justify-content: flex-start; }
  .hero-title { font-size: 2.35rem; }
  .footer-card { flex-direction: column; }
}
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


sample_candidates_path = REPO_ROOT / "uploads" / "sample_candidates.json"
sample_job_path = REPO_ROOT / "uploads" / "A1.txt"

DEFAULT_STATE = {
    "rank_rows": [],
    "rank_csv": "",
    "rank_errors": [],
    "last_run_summary": None,
    "candidate_details": {},
    "candidate_ready": False,
    "job_ready": False,
    "ranking_done": False,
    "validation_done": False,
    "download_done": False,
    "show_docs": False,
    "nav_choice": "Overview",
}
for key, value in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value


def icon_svg(name: str, size: int = 20) -> str:
    icons = {
        "target": '<circle cx="12" cy="12" r="10"/><circle cx="12" cy="12" r="6"/><circle cx="12" cy="12" r="2"/>',
        "upload": '<path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="17 8 12 3 7 8"/><line x1="12" y1="3" x2="12" y2="15"/>',
        "file": '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/>',
        "play": '<polygon points="6 3 20 12 6 21 6 3"/>',
        "download": '<path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/>',
        "check": '<polyline points="20 6 9 17 4 12"/>',
        "chart": '<line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/>',
        "book": '<path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M4 4v15.5A2.5 2.5 0 0 1 6.5 17H20V4z"/>',
        "github": '<path d="M9 19c-5 1.5-5-2.5-7-3m14 6v-3.9a3.4 3.4 0 0 0-1-2.6c3.3-.4 6.8-1.6 6.8-7A5.4 5.4 0 0 0 20 4.1 5 5 0 0 0 19.9 1S18.7.6 16 2.5a13.4 13.4 0 0 0-7 0C6.3.6 5.1 1 5.1 1A5 5 0 0 0 5 4.1a5.4 5.4 0 0 0-1.8 3.7c0 5.4 3.5 6.6 6.8 7a3.4 3.4 0 0 0-1 2.6V22"/>',
        "spark": '<path d="m12 3 1.9 5.8L20 11l-6.1 2.2L12 19l-1.9-5.8L4 11l6.1-2.2z"/>',
    }
    return (
        f'<svg width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" '
        f'stroke="currentColor" stroke-width="2.25" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
        f'{icons.get(name, icons["target"])}'
        f"</svg>"
    )


def format_bytes(size: int | float | None) -> str:
    if size is None:
        return "Unknown size"
    value = float(size)
    for unit in ("B", "KB", "MB", "GB"):
        if value < 1024 or unit == "GB":
            return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
        value /= 1024
    return f"{value:.1f} GB"


def uploaded_job_to_text(uploaded_file) -> str:
    """Read uploaded TXT/MD/DOCX job descriptions as plain text for the sandbox."""

    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - deployment dependency guard
            raise RuntimeError("DOCX support requires python-docx. Add python-docx to requirements.txt.") from exc

        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        return "\n".join(paragraphs + table_rows).strip()

    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


def count_candidates(path: Path) -> int | None:
    try:
        return sum(1 for _ in iter_candidates(path))
    except Exception:
        return None


def candidate_success_card(file_name: str, file_size: int | None, candidate_count: int | None) -> None:
    count_line = f"<br>Detected candidates: <strong>{candidate_count}</strong>" if candidate_count is not None else ""
    st.markdown(
        f"""
        <div class="success-card">
          <div class="success-icon">{icon_svg("check", 18)}</div>
          <div>
            <div class="success-title">Candidate sample uploaded</div>
            <div class="success-meta">
              {html.escape(file_name)} · {format_bytes(file_size)}{count_line}<br>
              Upload completed successfully.
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def job_success_card(file_name: str, file_size: int | None, source_note: str = "Upload completed successfully.") -> None:
    st.markdown(
        f"""
        <div class="success-card">
          <div class="success-icon">{icon_svg("check", 18)}</div>
          <div>
            <div class="success-title">Job description ready</div>
            <div class="success-meta">
              {html.escape(file_name)} · {format_bytes(file_size)}<br>
              {html.escape(source_note)}
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def load_candidate_details_for_ids(path: Path, candidate_ids: set[str]) -> dict[str, dict[str, Any]]:
    details: dict[str, dict[str, Any]] = {}
    if not candidate_ids:
        return details
    try:
        for candidate in iter_candidates(path):
            candidate_id = str(candidate.get("candidate_id") or "")
            if candidate_id not in candidate_ids:
                continue
            profile = candidate.get("profile") or {}
            skills = [str(item.get("name") or "") for item in candidate.get("skills") or [] if item.get("name")]
            details[candidate_id] = {
                "name": profile.get("anonymized_name") or candidate_id,
                "title": profile.get("current_title") or profile.get("headline") or "Candidate",
                "years": profile.get("years_of_experience"),
                "location": profile.get("location") or "",
                "skills": skills[:8],
                "summary": profile.get("summary") or "",
            }
            if len(details) == len(candidate_ids):
                break
    except Exception:
        return details
    return details


def extract_reason_skills(reasoning: str) -> list[str]:
    marker = "evidence includes "
    if marker not in reasoning:
        return []
    segment = reasoning.split(marker, 1)[1].split(";", 1)[0]
    if "few exact" in segment.lower():
        return []
    return [part.strip() for part in segment.split(",") if part.strip()][:5]


def candidate_display(row: Any) -> dict[str, Any]:
    details = st.session_state.candidate_details.get(row.candidate_id, {})
    skills = details.get("skills") or extract_reason_skills(row.reasoning)
    name = details.get("name") or row.candidate_id
    title = details.get("title") or "Candidate"
    years = details.get("years")
    years_text = f"{float(years):.1f} yrs" if isinstance(years, int | float) else "experience signal"
    summary = f"{title} with {years_text}. {row.reasoning}"
    return {"name": name, "title": title, "skills": skills[:6], "summary": summary}


def mark_downloaded() -> None:
    st.session_state.download_done = True


def checklist_item(index: int, label: str, done: bool) -> str:
    dot = "✓" if done else ""
    done_class = "done" if done else ""
    return (
        f'<div class="check-item">'
        f'<div class="check-dot {done_class}">{dot}</div>'
        f'<div class="check-label {done_class}">{index}. {html.escape(label)}</div>'
        f"</div>"
    )


def checklist_state() -> list[tuple[int, str, bool]]:
    return [
        (1, "Candidate sample", bool(st.session_state.candidate_ready)),
        (2, "Job description", bool(st.session_state.job_ready)),
        (3, "Run ranking", bool(st.session_state.ranking_done)),
        (4, "Validate final top-100 locally", bool(st.session_state.validation_done)),
        (5, "Download CSV", bool(st.session_state.download_done)),
    ]


def render_checklist() -> None:
    items = checklist_state()
    completed = sum(1 for _, _, done in items if done)
    percent = int(completed / len(items) * 100)
    st.markdown(
        f"""
        <div class="check-progress-label"><span>{completed} / {len(items)} Completed</span><span>{percent}%</span></div>
        <div class="check-progress-track"><div class="check-progress-fill" style="width:{percent}%"></div></div>
        <div class="checklist-wrap">{''.join(checklist_item(*item) for item in items)}</div>
        """,
        unsafe_allow_html=True,
    )


def status_html(label: str, value: str, icon_name: str) -> None:
    st.markdown(
        f"""
        <div class="status-card">
          <div class="status-icon">{icon_svg(icon_name, 28)}</div>
          <div class="status-label">{html.escape(label)}</div>
          <div class="status-value">{html.escape(value)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_topbar() -> None:
    st.markdown(
        f"""
        <div class="spotify-shell">
          <div class="topbar">
            <div class="brand-lockup">
              <div class="brand-icon">{icon_svg("target", 20)}</div>
              <div>
                <div class="brand-title">Redrob Candidate Ranker</div>
                <div class="brand-subtitle">Offline submission sandbox</div>
              </div>
            </div>
            <div class="topbar-badges">
              <span class="badge badge-green">{icon_svg("check", 16)} CPU-only</span>
              <span class="badge">No network</span>
              <span class="badge">DOCX ready</span>
              <span class="badge">CSV export</span>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_hero() -> None:
    st.markdown(
        """
        <div class="hero-card">
          <div class="hero-content">
            <h1 class="hero-title">Rank candidates with <span class="gradient-text">offline AI logic</span>.</h1>
            <p class="hero-subtitle">
              Upload candidates and a job description, run the deterministic Redrob ranker, validate top-100 output, and export CSV.
            </p>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def step_status(index: int, title: str, done: bool, active: bool) -> str:
    classes = "step-card"
    if done:
        classes += " completed"
    elif active:
        classes += " active"
    check = '<span class="step-check">✓</span>' if done else ""
    return (
        f'<div class="{classes}">{check}'
        f'<div class="step-kicker">STEP {index}</div>'
        f'<div class="step-title">{html.escape(title)}</div>'
        f"</div>"
    )


def active_step() -> int:
    if not st.session_state.candidate_ready:
        return 1
    if not st.session_state.job_ready:
        return 2
    if not st.session_state.ranking_done:
        return 3
    if not st.session_state.validation_done:
        return 4
    if not st.session_state.download_done:
        return 5
    return 5


def render_steps() -> None:
    active = active_step()
    step_defs = [
        (1, "Upload Candidates", st.session_state.candidate_ready),
        (2, "Upload Job Description", st.session_state.job_ready),
        (3, "Run Ranking", st.session_state.ranking_done),
        (4, "Validate Top-100", st.session_state.validation_done),
        (5, "Download CSV", st.session_state.download_done),
    ]
    html_steps = "".join(step_status(i, title, done, active == i and not done) for i, title, done in step_defs)
    st.markdown(f'<div class="step-grid">{html_steps}</div>', unsafe_allow_html=True)


def run_ranking(
    *,
    use_repo_sample: bool,
    candidate_upload: Any,
    use_repo_job: bool,
    job_upload: Any,
    job_text: str,
    top_k: int,
) -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)

        if use_repo_sample:
            candidates_path = sample_candidates_path
            candidate_source = "bundled sample_candidates.json"
        elif candidate_upload is not None:
            suffix = Path(candidate_upload.name).suffix or ".jsonl"
            candidates_path = tmp / f"candidates{suffix}"
            candidates_path.write_bytes(candidate_upload.getvalue())
            candidate_source = candidate_upload.name
        else:
            st.error("Please provide a candidate sample.")
            st.stop()

        if use_repo_job:
            job_path = sample_job_path
            job_source = "bundled A1.txt"
        elif job_upload is not None:
            try:
                extracted_job_text = uploaded_job_to_text(job_upload)
            except Exception as exc:
                st.error(f"Could not read job description file: {exc}")
                st.stop()
            if not extracted_job_text:
                st.error("The uploaded job description did not contain readable text.")
                st.stop()
            job_path = tmp / "job_description.txt"
            job_path.write_text(extracted_job_text, encoding="utf-8")
            job_source = job_upload.name
        elif job_text.strip():
            job_path = tmp / "job_description.txt"
            job_path.write_text(job_text, encoding="utf-8")
            job_source = "pasted text"
        else:
            st.error("Please provide a job description.")
            st.stop()

        progress = st.progress(0, text="Preparing files...")
        try:
            progress.progress(22, text="Reading candidate sample...")
            with st.spinner("Ranking candidates with deterministic offline scoring..."):
                progress.progress(58, text="Scoring candidate-job fit...")
                rows = rank_candidates(iter_candidates(candidates_path), top_k=top_k)
                progress.progress(84, text="Loading candidate preview data...")
                details = load_candidate_details_for_ids(candidates_path, {row.candidate_id for row in rows})
                progress.progress(92, text="Writing CSV...")
        except SystemExit as exc:
            progress.empty()
            st.error(str(exc))
            st.stop()
        except Exception as exc:
            progress.empty()
            st.exception(exc)
            st.stop()

        if not rows:
            progress.empty()
            st.error("No valid candidates found in the uploaded file.")
            st.stop()

        output_path = tmp / "sandbox_submission.csv"
        write_submission(rows, output_path)
        csv_text = output_path.read_text(encoding="utf-8")
        official_errors = validate_submission(str(output_path)) if len(rows) == 100 else []
        progress.progress(100, text="Ranking complete.")

        st.session_state.rank_rows = rows
        st.session_state.rank_csv = csv_text
        st.session_state.rank_errors = official_errors
        st.session_state.candidate_details = details
        st.session_state.last_run_summary = {
            "candidate_source": candidate_source,
            "job_source": job_source,
            "ranked_rows": len(rows),
            "top_score": rows[0].score,
            "top_candidate": rows[0].candidate_id,
        }
        st.session_state.candidate_ready = True
        st.session_state.job_ready = True
        st.session_state.ranking_done = True
        st.session_state.validation_done = len(rows) == 100 and not official_errors
        st.session_state.download_done = False


def score_class(score: float) -> str:
    if score >= 0.70:
        return "score-high"
    if score >= 0.40:
        return "score-mid"
    return "score-low"


def skill_chips(skills: list[str]) -> str:
    if not skills:
        return '<span class="skill-chip">Evidence in reason</span>'
    return "".join(f'<span class="skill-chip">{html.escape(skill)}</span>' for skill in skills[:5])


def render_preview_cards(rows: list[Any]) -> None:
    cards = []
    for index, row in enumerate(rows[:3], start=1):
        data = candidate_display(row)
        rank_class = " rank-one" if index == 1 else ""
        cards.append(
            f"""
            <div class="preview-card{rank_class}">
              <div class="preview-rank">Rank #{index}</div>
              <div class="preview-name">{html.escape(str(data['name']))}</div>
              <div class="preview-score">{row.score * 100:.1f}%</div>
              <div>{skill_chips(data['skills'])}</div>
              <p style="margin-top:12px; color:var(--text-2); font-size:0.9rem;">{html.escape(str(data['summary'])[:220])}</p>
            </div>
            """
        )
    st.markdown('<div class="preview-grid">' + "".join(cards) + "</div>", unsafe_allow_html=True)


def render_results_table(rows: list[Any]) -> None:
    body = []
    for index, row in enumerate(rows, start=1):
        data = candidate_display(row)
        top_class = " top-three" if index <= 3 else ""
        rank_one = " rank-one" if index == 1 else ""
        chip_class = "gold" if index == 1 else ""
        pct = max(2, min(100, row.score * 100))
        body.append(
            f"""
            <tr class="{top_class}{rank_one}">
              <td><span class="rank-chip {chip_class}">{index}</span></td>
              <td><div class="candidate-id">{html.escape(str(data['name']))}</div><div style="color:var(--muted); font-size:0.82rem;">{html.escape(row.candidate_id)}</div></td>
              <td><span class="score-pill {score_class(row.score)}">{row.score:.4f}</span><div class="match-track"><div class="match-fill" style="width:{pct:.1f}%"></div></div></td>
              <td>{pct:.1f}%</td>
              <td>{skill_chips(data['skills'])}</td>
              <td>{html.escape(row.reasoning)}</td>
            </tr>
            """
        )
    st.markdown(
        f"""
        <div class="results-table-wrap">
          <table class="results-table">
            <thead>
              <tr><th>Rank</th><th>Candidate</th><th>Match Score</th><th>Match %</th><th>Skills</th><th>Reason</th></tr>
            </thead>
            <tbody>{''.join(body)}</tbody>
          </table>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_results() -> None:
    rows = st.session_state.rank_rows
    if not rows:
        st.markdown(
            f"""
            <div class="empty-state">
              <div>
                <div class="empty-illustration">{icon_svg("chart", 42)}</div>
                <h3>No ranking generated yet.</h3>
                <p>Upload candidates and a job description to begin.</p>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        return

    st.success(f"Generated {len(rows)} ranked rows.")
    metric_cols = st.columns(4)
    metric_cols[0].metric("Ranked rows", len(rows))
    metric_cols[1].metric("Top score", f"{rows[0].score:.4f}")
    metric_cols[2].metric("Best candidate", rows[0].candidate_id)
    metric_cols[3].metric("CSV columns", "4")

    summary = st.session_state.last_run_summary or {}
    if summary:
        st.caption(f"Candidate source: {summary.get('candidate_source')} · Job source: {summary.get('job_source')}")

    if len(rows) == 100:
        if st.session_state.rank_errors:
            st.warning("CSV generated, but official validator found issues:")
            for error in st.session_state.rank_errors:
                st.write(f"- {error}")
        else:
            st.success("Official validator passed. Checklist item 4 is complete.")
    else:
        st.info("Official validator requires exactly 100 rows. Checklist item 4 completes only for a valid top-100 run.")

    st.download_button(
        "Download ranked CSV",
        data=st.session_state.rank_csv,
        file_name="sandbox_submission.csv",
        mime="text/csv",
        use_container_width=True,
        on_click=mark_downloaded,
    )

    st.subheader("Top 3 candidates")
    render_preview_cards(rows)

    st.subheader("Ranked candidates")
    render_results_table(rows)


with st.sidebar:
    st.markdown(
        f"""
        <div class="sidebar-brand">
          <div class="sidebar-logo">{icon_svg("target", 22)}</div>
          <div><div class="sidebar-title">Redrob Ranker</div></div>
        </div>
        <p class="sidebar-copy">Small-sample sandbox for the official offline ranking system.</p>
        """,
        unsafe_allow_html=True,
    )
    st.divider()

    page = st.radio(
        "Quick access",
        ["Overview", "Run Ranker", "Results"],
        label_visibility="collapsed",
        key="nav_choice",
    )

    st.divider()
    st.markdown("### Live checklist")
    render_checklist()
    st.divider()
    st.caption("CPU-only · no network · no hosted LLM calls")

render_topbar()
render_hero()

cta_cols = st.columns([1, 1, 4])
with cta_cols[0]:
    if st.button("Run Demo", use_container_width=True):
        st.session_state.nav_choice = "Run Ranker"
        st.rerun()
with cta_cols[1]:
    if st.button("Documentation", use_container_width=True):
        st.session_state.show_docs = not st.session_state.show_docs

status_cols = st.columns(4)
with status_cols[0]:
    status_html("Candidates", "Max 100", "target")
with status_cols[1]:
    status_html("Job Description", "TXT / MD / DOCX", "file")
with status_cols[2]:
    status_html("Ranking", "Offline Engine", "play")
with status_cols[3]:
    status_html("Output", "CSV Export", "download")

st.write("")
render_steps()

if page == "Overview":
    st.subheader("Getting started")
    st.markdown(
        """
        <div class="quick-card" style="max-width:760px;">
          <h3 class="card-title">Follow the live checklist to complete the ranking workflow.</h3>
          <p class="card-copy">Upload/select a candidate sample, add the job description, run ranking, validate a top-100 run, and download the CSV.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if st.session_state.show_docs:
        st.write("")
        st.subheader("Documentation")
        st.code(
            "python rank.py --candidates ./candidates.jsonl --job ./uploads/A1.txt --out ./team_yourid.csv --top-k 100\n"
            "python validate_submission.py ./team_yourid.csv",
            language="bash",
        )

elif page == "Run Ranker":
    left, right = st.columns(2)

    with left:
        st.markdown('<span class="step-pill">STEP 1 · Upload Candidates</span>', unsafe_allow_html=True)
        use_repo_sample = sample_candidates_path.exists() and st.checkbox("Use bundled sample_candidates.json", value=True)
        candidate_upload = None
        if not use_repo_sample:
            candidate_upload = st.file_uploader(
                "Upload candidate sample",
                type=["json", "jsonl", "ndjson", "txt", "gz"],
                help="Accepts JSON array, JSONL/NDJSON/TXT with one JSON object per line, or gzipped JSONL.",
            )
        st.session_state.candidate_ready = bool(use_repo_sample or candidate_upload is not None)
        if use_repo_sample:
            candidate_success_card("sample_candidates.json", sample_candidates_path.stat().st_size, count_candidates(sample_candidates_path))
        elif candidate_upload is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(candidate_upload.name).suffix or ".jsonl") as handle:
                handle.write(candidate_upload.getvalue())
                temp_candidate_path = Path(handle.name)
            candidate_success_card(candidate_upload.name, candidate_upload.size, count_candidates(temp_candidate_path))
            temp_candidate_path.unlink(missing_ok=True)
        st.caption("Sandbox samples can be ≤100 candidates. Use the CLI for the official 100K run.")

    with right:
        st.markdown('<span class="step-pill">STEP 2 · Upload Job Description</span>', unsafe_allow_html=True)
        use_repo_job = sample_job_path.exists() and st.checkbox("Use bundled A1.txt job description", value=True)
        job_upload = None
        job_text = ""
        if not use_repo_job:
            job_upload = st.file_uploader(
                "Upload job description",
                type=["txt", "md", "docx"],
                help="DOCX files are parsed with python-docx; text and markdown are read directly.",
            )
            job_text = st.text_area("Or paste job description", height=165, placeholder="Paste the job description here...")
        st.session_state.job_ready = bool(use_repo_job or job_upload is not None or job_text.strip())
        if use_repo_job:
            job_success_card("A1.txt", sample_job_path.stat().st_size)
        elif job_upload is not None:
            job_success_card(job_upload.name, job_upload.size)
        elif job_text.strip():
            job_success_card("Pasted job description", len(job_text.encode("utf-8")), "Text entered successfully.")
        st.caption("DOCX, TXT, MD, or pasted text are supported.")

    st.write("")
    top_k = st.slider(
        "Top K rows to generate",
        min_value=1,
        max_value=100,
        value=50 if use_repo_sample else 100,
        help="Official submission requires 100 rows. Small sandbox samples may contain fewer candidates.",
    )
    required_ready = bool(st.session_state.candidate_ready and st.session_state.job_ready)
    st.markdown('<div class="sticky-run-panel">', unsafe_allow_html=True)
    run_cols = st.columns([2, 1])
    with run_cols[0]:
        st.markdown(
            f"""
            <div class="card-title">STEP 3 · Run Ranking</div>
            <p class="card-copy">{'Ready to rank candidates.' if required_ready else 'Upload/select candidates and a job description to enable ranking.'}</p>
            """,
            unsafe_allow_html=True,
        )
    with run_cols[1]:
        run = st.button("Run Ranking", type="primary", use_container_width=True, disabled=not required_ready)
    st.markdown('</div>', unsafe_allow_html=True)

    if run:
        run_ranking(
            use_repo_sample=use_repo_sample,
            candidate_upload=candidate_upload,
            use_repo_job=use_repo_job,
            job_upload=job_upload,
            job_text=job_text,
            top_k=top_k,
        )
        render_results()

elif page == "Results":
    render_results()

st.markdown(
    """
    <div class="footer-card">
        <p style="margin:0; text-align:center; color:#9ca3af; font-size:14px;">
            Made with ❤️ by <strong style="color:white;">Core4</strong>
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)
